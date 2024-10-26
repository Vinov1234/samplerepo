'''import docx
doc = docx.Document('C:/Users/pc/Documents/files1.docx')
print(len(doc.paragraphs))
doc.paragraphs[0].text'''


import os
from docx import Document

# Check if the file exists
file_path = ("C:/Users/pc/Documents/files1.docx")
if os.path.exists(file_path):
    doc = Document(file_path)
    # Process the document
else:
    print("File not found. Please check the path.")
